# streamlit run app.py --server.port 8501 --server.address 0.0.0.0
from __future__ import annotations
import io
import json
from typing import List, Tuple, Dict, Any
from datetime import datetime
from embedded_house_prices_data import EMBEDDED_SAMPLE_CSV, build_dataset_from_embedded_csv

import numpy as np
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt

# ---- Import your toolkit (place the jmpkit/ folder next to this file) ----
from jmpkit import (
    Dataset, load_any, add_channel_expr, add_channel_func,
    jmp_distribution_report, jmp_multivariate_panel_full,
    fit_model, jmp_fit_and_plots, plot_xy_by_group,
    tag_fliers, get_fliers, set_flier_color, get_flier_color
)

# =========================================================
#            Performance helpers & cached funcs
# =========================================================

def _ensure_state():
    """Initialize frequently used session_state entries."""
    ss = st.session_state
    ss.setdefault("exports", {})
    ss.setdefault("dataset_info", "")
    ss.setdefault("data_version", 0)  # bump whenever data frame changes
    ss.setdefault("last_dist", {})
    ss.setdefault("tag_ui", {})

def _bump_data_version(n: int = 1):
    st.session_state["data_version"] = int(st.session_state.get("data_version", 0)) + int(n)

@st.cache_data(show_spinner=False)
def cached_summary(df_signature: Tuple[int, int, Tuple[str, ...], Tuple[str, ...]]) -> pd.DataFrame:
    """Cache the dataset summary. Signature should change when df changes."""
    # The actual DataFrame is read from session_state at call-time.
    return st.session_state["data"].summary()

@st.cache_data(show_spinner=False)
def to_csv_bytes(df_signature: Tuple[int, int, Tuple[str, ...]]) -> bytes:
    df = st.session_state["data"].df
    return df.to_csv(index=False).encode("utf-8")

@st.cache_data(show_spinner=False)
def to_xlsx_bytes(df_signature: Tuple[int, int, Tuple[str, ...]]) -> bytes:
    df = st.session_state["data"].df
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="data")
    bio.seek(0)
    return bio.read()

@st.cache_data(show_spinner=False)
def fig_to_png_bytes_cached(fig_id: str, dpi: int, version: int) -> bytes:
    """Cache fig->PNG bytes for export/render; fig is addressed indirectly by an id pattern.
    We generate bytes on demand when called, but Streamlit caches by (fig_id, dpi, version).
    The caller must pass the actual figure to _fig_to_png_bytes (uncached) when showing inline."""
    # This function is kept to allow future refactors to a fig registry if needed.
    # For now, it simply exists to keep a stable cache key space.
    return b""  # Not used directly; kept as placeholder for API stability

def _df() -> pd.DataFrame:
    return st.session_state["data"].df

def _df_signature() -> Tuple[int, int, Tuple[str, ...]]:
    df = _df()
    return (len(df), df.shape[1], tuple(df.columns))

def _df_signature_with_dtypes() -> Tuple[int, int, Tuple[str, ...], Tuple[str, ...]]:
    df = _df()
    return (len(df), df.shape[1], tuple(df.columns), tuple(map(str, df.dtypes)))

def _safe_num_cols(df: pd.DataFrame) -> List[str]:
    return list(df.select_dtypes(include=[np.number]).columns)

def _show_fig(fig, *, width_px: int, dpi: int, caption: str | None = None):
    """Render a Matplotlib figure to the page (no caching needed here)."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight")
    buf.seek(0)
    st.image(buf, caption=caption, width=width_px)
    buf.close()

def _fig_to_png_bytes(fig, *, dpi: int = 150) -> bytes:
    """Bytes for export; wrapped so we can place a cache in front if needed."""
    bio = io.BytesIO()
    fig.savefig(bio, format="png", dpi=dpi, bbox_inches="tight")
    bio.seek(0)
    return bio.read()

# =========================================================
#                     Word export helpers
# =========================================================

def _need_docx() -> Tuple[bool, Any]:
    try:
        import docx
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        return True, (docx, Inches, Pt, Cm, WD_ALIGN_PARAGRAPH, WD_ORIENT, qn, OxmlElement)
    except Exception as e:
        return False, e

def _init_doc(title: str, dataset_info: str | None, mod) -> Any:
    docx, Inches, Pt, Cm, WD_ALIGN_PARAGRAPH, WD_ORIENT, qn, OxmlElement = mod
    doc = docx.Document()

    # Margins
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(2.54)
        sec.right_margin = Cm(2.54)

    # Base font (Calibri 11)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(11)

    # Cover page
    doc.core_properties.title = title
    h = doc.add_paragraph()
    hr = h.add_run(title)
    hr.bold = True
    hr.font.size = Pt(20)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    srun = sub.add_run(datetime.now().strftime("%Y-%m-%d"))
    srun.italic = True
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if dataset_info:
        doc.add_paragraph("")  # spacer
        box = doc.add_table(rows=1, cols=1, style="Light Grid Accent 1")
        box.alignment = WD_ALIGN_PARAGRAPH.CENTER
        box.rows[0].cells[0].text = dataset_info

    # spacer
    doc.add_paragraph("")
    doc.add_paragraph("")

    return doc

def _add_caption(doc, text: str, mod):
    _, _, Pt, _, WD_ALIGN_PARAGRAPH, _, _, _ = mod
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)

def _add_section_heading(doc, text: str, level: int = 1):
    doc.add_heading(text, level=max(1, min(level, 4)))

def _repeat_header(table):
    """Ensure the first row repeats as header on page breaks."""
    tbl = table._element
    props = getattr(tbl, "tblPr", None)
    if props is None:
        props = tbl.get_or_add_tblPr()

    # Add <w:tblHeader w:val="true"/> if missing
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    has_tbl_header = props.xpath("./w:tblHeader")
    if not has_tbl_header:
        hdr = OxmlElement("w:tblHeader")
        hdr.set(qn("w:val"), "true")
        props.append(hdr)

def _add_df_to_docx_table(document, df: pd.DataFrame, title: str | None = None, *, style_name: str = "Light Grid Accent 1", max_rows: int = 1000):
    if title:
        _add_section_heading(document, title, level=3)

    if df is None or df.shape[0] == 0:
        p = document.add_paragraph("(no rows)")
        p.runs[0].italic = True
        return

    show_df = df.copy()
    truncated = False
    if len(show_df) > max_rows:
        show_df = show_df.iloc[:max_rows].copy()
        truncated = True

    show_df = show_df.map(lambda v: "" if pd.isna(v) else str(v))

    table = document.add_table(rows=1, cols=len(show_df.columns))
    try:
        table.style = style_name
    except Exception:
        pass

    hdr_cells = table.rows[0].cells
    for j, col in enumerate(show_df.columns):
        hdr_cells[j].text = str(col)

    for _, row in show_df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(show_df.columns):
            cells[j].text = row[col]

    _repeat_header(table)

    if truncated:
        note = document.add_paragraph()
        note_run = note.add_run(f"(Table truncated to first {max_rows:,} rows)")
        note_run.italic = True

def _add_image(document, img_bytes: bytes, caption: str | None, mod, width_inches: float = 6.5):
    docx, Inches, Pt, WD_CM, WD_ALIGN_PARAGRAPH, _, _, _ = mod
    document.add_picture(io.BytesIO(img_bytes), width=Inches(width_inches))
    if caption:
        _add_caption(document, caption, mod)

def _export_payload_to_docx(payload: Dict[str, Any], title: str) -> io.BytesIO:
    ok, mod = _need_docx()
    if not ok:
        st.error("Missing dependency `python-docx`. Install with: `pip install python-docx`")
        return io.BytesIO()

    doc = _init_doc(title, payload.get("dataset_info"), mod)

    # Data tab
    if payload.get("data_tab"):
        _add_section_heading(doc, "Data", level=2)
        dsec = payload["data_tab"]
        if dsec.get("preview_df") is not None:
            _add_df_to_docx_table(doc, dsec["preview_df"], "Preview (first rows)")
        if dsec.get("summary_df") is not None:
            _add_df_to_docx_table(doc, dsec["summary_df"], "Summary")

    # Distribution tab
    if payload.get("dist_tab"):
        _add_section_heading(doc, "Distribution", level=2)
        for item in payload["dist_tab"]:
            _add_section_heading(doc, str(item.get("column", "Column")), level=3)
            for cap, fig_bytes in item.get("figs", []):
                _add_image(doc, fig_bytes, cap, mod)
            for cap, df in item.get("tables", []):
                _add_df_to_docx_table(doc, df, cap)

    # Multivariate tab
    if payload.get("multi_tab"):
        _add_section_heading(doc, "Multivariate", level=2)
        m = payload["multi_tab"]
        for cap, fig_bytes in m.get("figs", []):
            _add_image(doc, fig_bytes, cap, mod)
        for cap, df in m.get("tables", []):
            _add_df_to_docx_table(doc, df, cap)
        if "UCL" in m:
            doc.add_paragraph(f"UCL = {m['UCL']:.3f}")

    # Fit Model tab
    if payload.get("fit_tab"):
        _add_section_heading(doc, "Fit Model", level=2)
        f = payload["fit_tab"]
        if f.get("persona"):
            doc.add_paragraph(f"Personality: {f['persona']}")
        if f.get("formula"):
            _add_section_heading(doc, "Formula", level=3)
            doc.add_paragraph(f["formula"])
        for sec_name, df in f.get("summary_tables", []):
            _add_df_to_docx_table(doc, df, f"Summary — {sec_name}")
        if f.get("metrics_json"):
            _add_section_heading(doc, "Metrics (JSON)", level=3)
            doc.add_paragraph(f["metrics_json"])
        for cap, fig_bytes in f.get("figs", []):
            _add_image(doc, fig_bytes, cap, mod)

    # XY Plot tab
    if payload.get("xy_tab"):
        _add_section_heading(doc, "Fit Y by X (faceted)", level=2)
        xyt = payload["xy_tab"]
        if xyt.get("desc"):
            doc.add_paragraph(xyt["desc"])
        for cap, fig_bytes in xyt.get("figs", []):
            _add_image(doc, fig_bytes, cap, mod)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

@st.cache_data(show_spinner=False)
def export_payload_to_docx_cached(payload: Dict[str, Any], title: str, version: int) -> bytes:
    """Cache the final .docx bytes by payload 'shape' (we rely on caller to pass stable payload)."""
    bio = _export_payload_to_docx(payload, title)
    return bio.read()

def _export_button_for_tab(tab_key: str, file_label: str = "Export this tab to Word"):
    exports = st.session_state.setdefault("exports", {})
    tab_payload = exports.get(tab_key)
    if not tab_payload:
        st.caption("Nothing to export yet on this tab.")
        return
    title = f"JMP-like Analysis — {tab_key.replace('_',' ').title()}"
    payload = {"dataset_info": st.session_state.get("dataset_info", ""), **{tab_key: tab_payload}}
    # Cache by a cheap signature: tab_key + data_version + simple sizes
    version = st.session_state["data_version"]
    try:
        docx_bytes = export_payload_to_docx_cached(payload, title, version)
    except Exception:
        # Fallback uncached if something went wrong with hashing inside cache
        docx_bytes = _export_payload_to_docx(payload, title).read()

    if not docx_bytes:
        return
    st.download_button(
        file_label,
        data=docx_bytes,
        file_name=f"{tab_key}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key=f"dl_{tab_key}"
    )

# =========================================================
#                       Data loading
# =========================================================

def _load_uploaded(file) -> Dataset:
    """Load uploaded file (Excel/CSV/Parquet/Feather/JSON/PKL) into Dataset."""
    name = file.name.lower()
    if name.endswith((".xlsx", ".xls")):
        df = pd.read_excel(file, sheet_name=0, header=0)
    elif name.endswith((".csv", ".tsv", ".txt")):
        df = pd.read_csv(file, sep=None, engine="python", header=0)
    elif name.endswith(".parquet"):
        df = pd.read_parquet(file)
    elif name.endswith(".feather"):
        df = pd.read_feather(file)
    elif name.endswith(".json"):
        try:
            df = pd.read_json(file, orient="records")
        except ValueError:
            file.seek(0)
            df = pd.read_json(file, orient="table")
    elif name.endswith(".pkl"):
        df = pd.read_pickle(file)
    else:
        st.error(f"Unsupported file type: {name}")
        st.stop()
    return Dataset(df.convert_dtypes(), source=name)

# =========================================================
#                       App Layout
# =========================================================

st.set_page_config(page_title="JMP-like Analysis — Web UI", layout="wide")
_ensure_state()

st.title("JMP-like Analysis — Web Interface")

with st.sidebar:
    st.header("1) Upload Data")
    up = st.file_uploader(
        "Excel/CSV/Parquet/Feather/JSON/PKL",
        type=["xlsx","xls","csv","tsv","txt","parquet","feather","json","pkl"],
        key="uploader_main"
    )
    if up is not None:
        st.session_state["data"] = _load_uploaded(up)
        st.session_state["dataset_info"] = f"Source: {up.name} — {len(_df()):,} rows × {_df().shape[1]} cols"
        _bump_data_version()
        st.success(f"Loaded: {up.name}  →  {len(_df()):,} rows × {_df().shape[1]} cols")

    if st.button("Load sample data (embedded CSV)", key="btn_load_sample_embedded_csv"):
        try:
            ds = build_dataset_from_embedded_csv(Dataset)
            st.session_state["data"] = ds
            st.session_state["dataset_info"] = (
                f"Source: house_prices_US.csv (embedded) — "
                f"{len(_df()):,} rows × {_df().shape[1]} cols"
            )
            _bump_data_version()
            st.success(
                f"Loaded sample (embedded CSV): house_prices_US  →  "
                f"{len(_df()):,} rows × {_df().shape[1]} cols"
            )
        except Exception as e:
            st.error(f"Failed to load embedded CSV sample: {e}")

    if "data" in st.session_state:
        st.divider()
        st.header("2) Channels")
        st.caption("Add derived columns using expressions. Use backticks for names with spaces: `temp(f)`")

        # ---- FORM to stop per-keystroke reruns ----
        with st.form("channels_form"):
            col1, col2 = st.columns([1,2])
            with col1:
                new_name = st.text_input("New column name", key="channel_name")
            with col2:
                expr = st.text_input("Expression (pandas.eval)", placeholder="e.g., log(`credit_score` + 1)", key="channel_expr")
            submitted_add = st.form_submit_button("Add Channel", disabled=not new_name or not expr)

        if submitted_add:
            try:
                add_channel_expr(st.session_state["data"], new_name, expr, overwrite=True)
                _bump_data_version()
                st.success(f"Added/updated channel: {new_name}")
            except Exception as e:
                st.error(f"Failed to add channel: {e}")

        st.divider()
        st.header("3) Display Size (no changes to jmpkit)")
        width_px = st.slider("Plot width (pixels)", 300, 2000, 900, 50, key="plot_width")
        dpi      = st.slider("Render DPI", 72, 300, 120, 4, key="plot_dpi")
        st.caption("These only affect how figures are rendered here; the underlying figure objects are unchanged.")

        st.divider()
        st.header("4) Download (cached)")
        if "data" in st.session_state:
            sig = _df_signature()
            st.download_button(
                "Download Excel (.xlsx)",
                data=to_xlsx_bytes(sig),
                file_name="dataset.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="dl_xlsx_main"
            )
            st.download_button(
                "Download CSV",
                data=to_csv_bytes(sig),
                file_name="dataset.csv",
                mime="text/csv",
                key="dl_csv_main"
            )

# Halt early if no data
if "data" not in st.session_state:
    st.info("Upload a file to get started. Once loaded, you can add channels and run analyses.")
    st.stop()

# =========================================================
#                         Tabs
# =========================================================

tab_data, tab_dist, tab_multi, tab_fit, tab_plot_scatter = st.tabs(
    ["Data", "Distribution", "Multivariate", "Fit Model", "Fit Y by X (faceted)"]
)

# ------------------ Data tab ------------------
with tab_data:
    st.subheader("Preview")
    prev_df = _df().head(100)
    st.dataframe(prev_df, width='stretch')

    st.subheader("Summary")
    sig_full = _df_signature_with_dtypes()
    summary_df = cached_summary(sig_full)
    st.dataframe(summary_df, width='stretch')

    # Capture for export
    st.session_state["exports"]["data_tab"] = {
        "preview_df": prev_df,
        "summary_df": summary_df,
    }

    st.markdown("---")
    _export_button_for_tab("data_tab", "Export this tab to Word")

# ------------------ Distribution tab ------------------
with tab_dist:
    st.subheader("Distribution Report")
    cols_sel = st.multiselect(
        "Columns (numeric suggested)",
        options=list(_df().columns),
        default=_safe_num_cols(_df())[:1],
        key="dist_cols_sel"
    )

    # container to accumulate export items for this run
    dist_items_for_export: List[Dict[str, Any]] = []

    if st.button("Run Distribution", disabled=len(cols_sel) == 0, key="btn_run_dist"):
        st.session_state["last_dist"].clear()
        for c in cols_sel:
            st.markdown(f"#### {c}")
            try:
                res = jmp_distribution_report(_df(), c)
                _show_fig(res["fig_hist"], width_px=st.session_state.get("plot_width", 900), dpi=st.session_state.get("plot_dpi", 120), caption="Histogram + Fitted Normal")
                _show_fig(res["fig_box"],  width_px=st.session_state.get("plot_width", 900), dpi=st.session_state.get("plot_dpi", 120), caption="Box Plot")
                st.dataframe(res["quantiles_df"], width='stretch')
                st.dataframe(res["summary_df"], width='stretch')
                st.dataframe(res["fit_params_df"], width='stretch')
                st.dataframe(res["fit_stats_df"], width='stretch')
                st.dataframe(res["gof_summary_df"], width='stretch')
                st.dataframe(res["ad_details_df"], width='stretch')

                # remember for tagging
                st.session_state["last_dist"][c] = {
                    "flier_index": res.get("flier_index", []),
                    "bounds": res.get("flier_bounds", (None, None)),
                }

                # collect for export
                figs = [
                    ("Histogram + Fitted Normal", _fig_to_png_bytes(res["fig_hist"], dpi=st.session_state.get("plot_dpi", 120))),
                    ("Box Plot", _fig_to_png_bytes(res["fig_box"],  dpi=st.session_state.get("plot_dpi", 120))),
                ]
                tables = [
                    ("Quantiles", res["quantiles_df"]),
                    ("Summary", res["summary_df"]),
                    ("Fit Parameters", res["fit_params_df"]),
                    ("Fit Stats", res["fit_stats_df"]),
                    ("GoF Summary", res["gof_summary_df"]),
                    ("AD Details", res["ad_details_df"]),
                ]
                dist_items_for_export.append({"column": c, "figs": figs, "tables": tables})

            except Exception as e:
                st.warning(f"Skipping {c}: {e}")

        # store export payload for this tab
        st.session_state["exports"]["dist_tab"] = dist_items_for_export

    # ---- Tagging controls (after results) ----
    if st.session_state["last_dist"]:
        st.markdown("---")
        st.markdown("### Tagging Options for Selected Channels")
        with st.form("tagging_form", clear_on_submit=False):
            for c in cols_sel:
                default_color = get_flier_color(c)
                default_checked = len(get_fliers(c)) > 0
                ui_state = st.session_state["tag_ui"].setdefault(c, {"checked": default_checked, "color": default_color})

                colA, colB, colC, colD = st.columns([2, 2, 2, 6])
                with colA:
                    st.write(f"**{c}**")
                with colB:
                    ui_state["checked"] = st.checkbox("Tag fliers", value=ui_state["checked"], key=f"tag_{c}")
                with colC:
                    ui_state["color"] = st.color_picker("Color", value=ui_state["color"], key=f"color_{c}")
                with colD:
                    n_new = len(st.session_state["last_dist"].get(c, {}).get("flier_index", []))
                    n_tag = len(get_fliers(c))
                    st.caption(f"New detected: **{n_new}** — currently tagged: **{n_tag}**")
            submitted = st.form_submit_button("Apply Tags")

        if submitted:
            applied_any = False
            for c in cols_sel:
                ui = st.session_state["tag_ui"][c]
                set_flier_color(c, ui["color"])
                if ui["checked"]:
                    idxs = st.session_state["last_dist"].get(c, {}).get("flier_index", [])
                    tag_fliers(c, idxs, mode="set")
                    applied_any = True
                else:
                    tag_fliers(c, [], mode="clear")

            if applied_any:
                st.success("Tags/colors applied. Plots will reflect these choices.")
            else:
                st.info("No channels tagged. (Any previous tags were cleared if unchecked.)")

    st.markdown("---")
    _export_button_for_tab("dist_tab", "Export this tab to Word")

# ------------------ Multivariate tab ------------------
with tab_multi:
    st.subheader("Multivariate Panel")
    mcols = st.multiselect("Columns (≥2)", options=list(_df().columns), default=_safe_num_cols(_df())[:3], key="multi_cols_sel")
    alpha = st.number_input("Alpha (tail prob for UCL)", min_value=0.0001, max_value=0.2, value=0.05, step=0.001, format="%.3f", key="multi_alpha")
    upper = st.selectbox("Upper triangle", options=["scatter", "corr"], index=0, key="multi_upper")
    diag = st.selectbox("Diagonal", options=["label", "none"], index=0, key="multi_diag")

    if st.button("Run Multivariate", disabled=len(mcols) < 2, key="btn_run_multi"):
        try:
            out = jmp_multivariate_panel_full(_df(), columns=mcols, alpha=alpha, upper=upper, diag=diag)
            _show_fig(out["fig_matrix"], width_px=st.session_state.get("plot_width", 900), dpi=st.session_state.get("plot_dpi", 120), caption="Scatterplot Matrix")
            st.dataframe(out["corr_df"], width='stretch')
            st.dataframe(out["n_pairs_df"], width='stretch')
            _show_fig(out["md_fig"], width_px=st.session_state.get("plot_width", 900), dpi=st.session_state.get("plot_dpi", 120), caption="Mahalanobis Distances")

            # capture for export
            figs = [
                ("Scatterplot Matrix", _fig_to_png_bytes(out["fig_matrix"], dpi=st.session_state.get("plot_dpi", 120))),
                ("Mahalanobis Distances", _fig_to_png_bytes(out["md_fig"], dpi=st.session_state.get("plot_dpi", 120))),
            ]
            tables = [
                ("Correlation Matrix", out["corr_df"]),
                ("N Pairs", out["n_pairs_df"]),
            ]
            st.session_state["exports"]["multi_tab"] = {"figs": figs, "tables": tables, "UCL": out.get("UCL")}
        except Exception as e:
            st.error(f"Multivariate failed: {e}")

    st.markdown("---")
    _export_button_for_tab("multi_tab", "Export this tab to Word")

# ------------------ Fit Model tab ------------------
with tab_fit:
    st.subheader("Fit Model")
    persona = st.selectbox(
        "Personality",
        [
            "standard_least_squares","stepwise","glm","generalized_regression",
            "mixed_model","manova","nominal_logistic","ordinal_logistic",
            "proportional_hazards","parametric_survival","partial_least_squares","response_screening"
        ],
        index=0,
        key="fit_persona"
    )
    y_all = list(_df().columns)
    eff_all = list(_df().columns)
    multi_y = persona in ("manova","partial_least_squares","response_screening")

    y_sel = st.multiselect("Y variables" if multi_y else "Y variable",
                           options=y_all, default=(y_all[:2] if multi_y else y_all[:1]),
                           key="fit_y_sel")
    if not multi_y and y_sel:
        y_sel = y_sel[0]

    effects = st.multiselect("Construct Model Effects", options=eff_all, default=eff_all[:2], key="fit_effects")

    colA, colB, colC = st.columns(3)
    with colA:
        degree = st.number_input("Polynomial degree (per effect)", 1, 5, 1, 1, key="fit_degree")
    with colB:
        cross = st.checkbox("Include pairwise interactions", value=False, key="fit_cross")
    with colC:
        add_poly = st.checkbox("Add polynomial terms", value=False, key="fit_add_poly")

    extra_kwargs = {}
    
    if persona == "stepwise":
        # --- JMP-like Stepwise Regression Control ---
        st.markdown("### Stepwise Regression Control")

        # Stopping rule (Criterion) and probabilities
        col0, col1, col2, col3 = st.columns([1,1,1,1])
        with col0:
            stopping_rule = st.selectbox("Stopping Rule", ["P-value Threshold","AIC"], index=0, key="sw_rule")
            extra_kwargs["selection_criterion"] = "pvalue" if stopping_rule == "P-value Threshold" else "aic"
        with col1:
            extra_kwargs["entry_p"] = st.number_input("Prob to Enter", min_value=0.0001, max_value=0.9999, value=float(st.session_state.get("fit_step_entryp_val", 0.05)), step=0.005, format="%.4f", key="sw_entry")
        with col2:
            extra_kwargs["exit_p"]  = st.number_input("Prob to Leave",  min_value=0.0001, max_value=0.9999, value=float(st.session_state.get("fit_step_exitp_val", 0.10)), step=0.005, format="%.4f", key="sw_exit")
        with col3:
            extra_kwargs["stepwise_direction"] = st.selectbox("Direction", ["Forward","Backward","Mixed","Both"], index=2, key="sw_dir").lower()

        extra_kwargs["max_steps"] = st.number_input("Max steps", 1, 500, 50, 1, key="sw_max_steps")

        # Maintain stepwise session state (current selected terms)
        ss = st.session_state
        if "sw_selected" not in ss:
            ss["sw_selected"] = []
        if "sw_running" not in ss:
            ss["sw_running"] = False

        # The available effects are the 'effects' list chosen above
        # Buttons row: Enter All, Remove All, Make Model, Run Model
        cA, cB, cC, cD = st.columns(4)
        with cA:
            if st.button("Enter All", key="sw_enter_all"):
                ss["sw_selected"] = list(effects)
        with cB:
            if st.button("Remove All", key="sw_remove_all"):
                ss["sw_selected"] = []
        with cC:
            make_model = st.button("Make Model", key="sw_make_model")
        with cD:
            run_model = st.button("Run Model", key="sw_run_model")

        # Row: Go, Stop, Step
        cE, cF, cG = st.columns(3)
        with cE:
            go = st.button("Go", key="sw_go")
        with cF:
            stop = st.button("Stop", key="sw_stop")
        with cG:
            step_once = st.button("Step", key="sw_step")

        # Compute actions
        def _fit_stepwise(initial_selected, max_steps):
            return fit_model(_df(), y=y_sel, effects=effects, personality="stepwise",
                             initial_selected=initial_selected,
                             degree=degree, cross=cross, add_poly=add_poly, **extra_kwargs, max_steps=max_steps)

        sw_path = []
        sw_selected = list(ss["sw_selected"])
        if make_model:
            # Fit OLS using the currently selected set
            if len(sw_selected) == 0:
                st.warning("No terms selected. Add terms with Enter All or Stepwise selection.")
            else:
                rhs = " + ".join(sw_selected)
                res = fit_model(_df(), y=y_sel, effects=sw_selected, personality="standard_least_squares",
                                degree=1, cross=False, add_poly=False)
                st.success(f"Made model: {y_sel} ~ {rhs}")
                st.code(res.formula or f"{y_sel} ~ {rhs}", language="text")
        if run_model or go:
            # Full run to stopping rule
            res = _fit_stepwise(sw_selected, max_steps=int(extra_kwargs["max_steps"]))
            m = res.metrics.get(y_sel, {})
            ss["sw_selected"] = list(m.get("selected_terms", sw_selected))
            sw_path = m.get("path", [])
        if step_once:
            # Run exactly one step from current state
            res = _fit_stepwise(sw_selected, max_steps=1)
            m = res.metrics.get(y_sel, {})
            ss["sw_selected"] = list(m.get("selected_terms", sw_selected))
            sw_path = m.get("path", [])

        # Display current selection and (latest) path increment
        st.markdown("**Current Selected Terms:** " + (", ".join(ss["sw_selected"]) if ss["sw_selected"] else "(none)"))
        if sw_path:
            st.markdown("**Last Step(s):**")
            st.dataframe(pd.DataFrame(sw_path), use_container_width=True)
        st.caption("Use Step to iterate one action at a time, or Go/Run Model to complete the procedure.")
    if persona == "glm":
        col1, col2 = st.columns(2)
        with col1:
            family = st.selectbox("Family", ["gaussian","binomial","poisson","gamma","inverse_gaussian","nb"], index=0, key="fit_glm_family")
            extra_kwargs["family"] = family
        with col2:
            link = st.selectbox("Link (optional)", ["", "identity","log","logit","probit","cloglog"], index=0, key="fit_glm_link")
            extra_kwargs["link"] = (None if link == "" else link)
    if persona == "mixed_model":
        if len(eff_all) > 0:
            extra_kwargs["groups"] = st.selectbox("Grouping column (random intercept)", eff_all, key="fit_mixed_group")
    if persona in ("proportional_hazards","parametric_survival"):
        dur = st.selectbox("Duration column", options=eff_all, key="fit_surv_dur")
        evt = st.selectbox("Event column (1=event)", options=eff_all, key="fit_surv_evt")
        extra_kwargs["duration_col"] = dur
        extra_kwargs["event_col"] = evt
        if persona == "parametric_survival":
            extra_kwargs["aft_distribution"] = st.selectbox("AFT distribution", ["weibull","lognormal","loglogistic","exponential"], index=0, key="fit_aft_dist")
    if persona == "partial_least_squares":
        extra_kwargs["n_components"] = st.number_input("PLS components", 1, 10, 2, 1, key="fit_pls_comp")

    if st.button("Fit", key="btn_fit"):
        try:
            result = fit_model(_df(), y=y_sel, effects=effects, personality=persona, degree=degree, cross=cross, add_poly=add_poly, **extra_kwargs)
            st.success(f"Fitted: {result.personality}")
            st.code(result.formula or "(no formula string)", language="text")

            # export capture: summaries as tables if possible
            summary_tables: List[Tuple[str, pd.DataFrame]] = []
            for k, v in result.summaries.items():
                st.markdown(f"#### Summary — {k}")
                if isinstance(v, dict):
                    for section, content in v.items():
                        st.markdown(f"**{section}**")
                        if isinstance(content, list) and content and isinstance(content[0], dict):
                            df = pd.DataFrame(content)
                            if "Term" in df.columns: df = df.set_index("Term")
                            st.dataframe(df, width='stretch')
                            summary_tables.append((f"{k} — {section}", df))
                        elif isinstance(content, dict):
                            first_val = next(iter(content.values()), {})
                            if isinstance(first_val, dict):
                                df = pd.DataFrame(content).T
                            else:
                                df = pd.DataFrame([content])
                            st.dataframe(df, width='stretch')
                            summary_tables.append((f"{k} — {section}", df))
                        else:
                            st.code(str(content), language="text")
                    continue
                try:
                    html = v.as_html()
                    st.code(html, height=520, scrolling=True)
                except Exception:
                    st.code(str(v), language="text")

            # metrics pretty JSON for export
            metrics_json = None
            if result.metrics:
                st.markdown("#### Metrics")
                try:
                    pretty = {k: {m: (float(vv) if hasattr(vv, "__float__") else vv) for m, vv in d.items()} for k, d in result.metrics.items()}
                    metrics_json = json.dumps(pretty, indent=2)
                    st.code(metrics_json, language="json")
                except Exception:
                    st.code(str(result.metrics), language="text")
                    metrics_json = str(result.metrics)

            figs_export: List[Tuple[str, bytes]] = []
            if persona in ("standard_least_squares","glm","generalized_regression","stepwise") and (not multi_y):
                st.markdown("#### Diagnostic Plots")
                regressor = effects[0] if effects else None
                figs = jmp_fit_and_plots(_df(), y=y_sel if isinstance(y_sel, str) else y_sel[0], effects=effects, personality=persona, regressor=regressor)
                one = next(iter(figs.values()))
                for fig in one["fig_regression"]:
                    _show_fig(fig, width_px=st.session_state.get("plot_width", 900), dpi=st.session_state.get("plot_dpi", 120), caption="Regression Plot")
                    figs_export.append(("Regression Plot", _fig_to_png_bytes(fig, dpi=st.session_state.get("plot_dpi", 120))))
                _show_fig(one["fig_actual"], width_px=st.session_state.get("plot_width", 900), dpi=st.session_state.get("plot_dpi", 120), caption="Actual by Predicted")
                figs_export.append(("Actual by Predicted", _fig_to_png_bytes(one["fig_actual"], dpi=st.session_state.get("plot_dpi", 120))))
                _show_fig(one["fig_residual"], width_px=st.session_state.get("plot_width", 900), dpi=st.session_state.get("plot_dpi", 120), caption="Residual by Predicted")
                figs_export.append(("Residual by Predicted", _fig_to_png_bytes(one["fig_residual"], dpi=st.session_state.get("plot_dpi", 120))))

            st.session_state["exports"]["fit_tab"] = {
                "persona": result.personality,
                "formula": result.formula or "",
                "summary_tables": summary_tables,
                "metrics_json": metrics_json,
                "figs": figs_export,
            }

        except Exception as e:
            st.error(f"Fit failed: {e}")

    st.markdown("---")
    _export_button_for_tab("fit_tab", "Export this tab to Word")

# ------------------ Plot Scatter tab ------------------
with tab_plot_scatter:
    st.subheader("Fit Y by X (faceted)")
    cols_all = list(_df().columns)
    x = st.selectbox("X (predictor)", options=cols_all, index=0 if cols_all else None, key="xy_x")
    y_opts = [c for c in cols_all if c != x] if x is not None else []
    y = st.selectbox("Y (response)", options=y_opts, index=0 if y_opts else None, key="xy_y")

    group = st.selectbox("Group (optional)", options=["(none)"] + cols_all, index=0, key="xy_group")
    grp = None if group == "(none)" else group

    ncols = st.slider("Facets per row", 1, 4, 3, key="xy_ncols")
    h = st.slider("Facet height (inches)", 3.0, 6.0, 4.0, key="xy_h")
    w = st.slider("Facet width (inches)", 3.0, 6.0, 4.0, key="xy_w")

    run_disabled = (x is None) or (y is None)
    if st.button("Run XY Plot", disabled=run_disabled, key="btn_run_xy"):
        fig, _ = plot_xy_by_group(_df(), x=x, y=y, group=grp, ncols=ncols, height=h, width=w)
        st.pyplot(fig, width='stretch')
        st.session_state["exports"]["xy_tab"] = {
            "desc": f"Y: {y} by X: {x}" + ("" if grp is None else f" — grouped by {grp}"),
            "figs": [("Faceted Scatter", _fig_to_png_bytes(fig, dpi=st.session_state.get("plot_dpi", 120)))],
        }

    st.markdown("---")
    _export_button_for_tab("xy_tab", "Export this tab to Word")
